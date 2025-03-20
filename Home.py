import streamlit as st
import pandas as pd
import fitz  # PyMuPDF for PDFs
import docx
import spacy
import numpy as np
import pytesseract
from PIL import Image
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
from docx import Document
from fpdf import FPDF
import io

# Load NLP models
nlp = spacy.load("en_core_web_sm")
bert_model = SentenceTransformer("all-MiniLM-L6-v2")

# Function to extract text from PDF, including image-based PDFs
def extract_text_from_pdf(uploaded_file):
    text = ""
    try:
        with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
            for page in doc:
                page_text = page.get_text("text")
                if not page_text.strip():  # If no text is extracted, use OCR
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    page_text = pytesseract.image_to_string(img)
                text += page_text + "\n"
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
    return text.strip()

# Function to extract text from Docx
def extract_text_from_docx(uploaded_file):
    text = ""
    try:
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
    return text.strip()

# Function to compute TF-IDF cosine similarity
def compute_tfidf_similarity(jd_text, resume_text):
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform([jd_text, resume_text])
    cosine_sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:])
    return cosine_sim.flatten()[0]

# Function to compute BERT similarity
def compute_bert_similarity(jd_text, resume_text):
    jd_embedding = bert_model.encode([jd_text], convert_to_tensor=True).cpu().numpy()
    resume_embedding = bert_model.encode([resume_text], convert_to_tensor=True).cpu().numpy()
    similarity = cosine_similarity(jd_embedding, resume_embedding)
    return similarity.flatten()[0]

# Function to normalize scores safely
def normalize_scores(scores):
    if len(scores) == 0:
        return np.array([])
    min_score, max_score = np.min(scores), np.max(scores)
    if max_score == min_score:  # Avoid division by zero
        return np.ones(len(scores)) * 50  # Assign a neutral score for all
    normalized = ((scores - min_score) / (max_score - min_score)) * 90 + 10  # Scale to 10-100
    return normalized

# Function to convert DOCX to PDF
def convert_docx_to_pdf(docx_file):
    doc = Document(docx_file)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    for para in doc.paragraphs:
        pdf.multi_cell(0, 10, para.text)

    pdf_output = io.BytesIO()
    pdf.output(pdf_output, "F")
    return pdf_output.getvalue()

# Streamlit UI
st.set_page_config(page_title="AI Resume Screening", layout="wide")
st.title("AI Resume Screening & Ranking System")

# Add custom CSS for styling
st.markdown(
    """
    <style>
    /* Main page styling */
    body {
        background-color: #f0f2f6;
        font-family: Arial, sans-serif;
    }
    .stTextArea textarea {
        font-size: 16px !important;
        height: 150px !important;
        width: 100% !important;
        padding: 10px;
        border-radius: 8px;
        border: 1px solid #ccc;
    }
    .stButton button {
        background-color: none;
        color:  #5c5c5c;
        padding: 7px 10px;
        font-size: 25px;
        cursor: pointer;
        text-align: left;
        border = none;
        
    }
    .stButton button:hover {
        background-color: #999999;
        color: black;
    }
    .stDataFrame {
        border-radius: 8px;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    /* Sidebar styling */
    .css-1d391kg {
        background-color: #ffffff;
        padding: 10px;
        border-radius: 8px;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    .sidebar .stButton button {
        width: 100%;
        margin-bottom: 10px;
    }
    /* About section styling */
    .about-container {
        display: flex;
        align-items: center;
        font-size: 24px;
        font-weight: 700;
        color: #5c5c5c;
        margin-bottom: 20px;
    }
    .stImage image{
        width = 150px;
        height = 150px;
        border-radius = 50%;
        object-fit - cover;
    }
    .info-icon {
        display: flex;
        justify-content: center;
        align-items: center;
        font-size: 20px;
        font-weight: bold;
        width: 30px;
        height: 30px;
        border-radius: 50%;
        background-color: white;
        color: #999999;
        margin-right: 10px;
        border: 3px solid #999999;
        margin-left: 10px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Initialize session state for resume data if not present
if "resume_texts" not in st.session_state:
    st.session_state.resume_texts = {}

# Sidebar UI
with st.sidebar:
    st.image(r"D:\al\project1\Resume_Ranking\images\shutterstock_546995980.webp")
    st.header("AI Resume Screening & Ranking")
    st.write(" ")
    st.write(" ")
    st.write(" ")
    st.write(" ")
    st.write(" ")
    st.write(" ")
    st.write(" ")
    st.write(" ")
    st.write(" ")
    st.write(" ")
    st.write(" ")
    st.write(" ")
    st.write(" ")
    st.write(" ")

    st.header("Settings")

    # Search Resumes Button
    if st.button("Search Resumes"):
        st.session_state["show_search"] = True
        st.session_state["show_about"] = False

    # About Button
    if st.button("About"):
        st.session_state["show_about"] = True
        st.session_state["show_search"] = False

    # Download Results Button
    if st.button("Download Results"):
        if "results_df" in st.session_state and not st.session_state["results_df"].empty:
            csv = st.session_state["results_df"].to_csv(index=False).encode('utf-8')
            st.download_button(
                label="Download Ranked Resumes",
                data=csv,
                file_name="ranked_resumes.csv",
                mime="text/csv"
            )
        else:
            st.warning("No results available to download.")

# Job description input
jd_text = st.text_area("Enter Job Description:").strip()
uploaded_resumes = st.file_uploader("Upload Resumes (PDF/DOCX)", type=["pdf", "docx"], accept_multiple_files=True)

if jd_text and uploaded_resumes:
    resume_texts = {}
    resume_names = []
    
    for resume in uploaded_resumes:
        if resume.name in st.session_state.resume_texts:
            text = st.session_state.resume_texts[resume.name]
        else:
            text = ""
            if resume.name.endswith(".pdf"):
                text = extract_text_from_pdf(resume)
            elif resume.name.endswith(".docx"):
                text = extract_text_from_docx(resume)
            st.session_state.resume_texts[resume.name] = text if text.strip() else "N/A"
        
        if text.strip() != "N/A":
            resume_texts[resume.name] = text
    
    if resume_texts:
        resume_names = list(resume_texts.keys())
        resume_text_values = list(resume_texts.values())

        tfidf_scores = np.array([compute_tfidf_similarity(jd_text, rt) for rt in resume_text_values])
        bert_scores = np.array([compute_bert_similarity(jd_text, rt) for rt in resume_text_values])
        
        tfidf_scores = normalize_scores(tfidf_scores)
        bert_scores = normalize_scores(bert_scores)
        
        results_df = pd.DataFrame({
            "Resume Name": resume_names,
            "TF-IDF Score": tfidf_scores,
            "BERT Score": bert_scores
        })
        
        results_df["Final Score"] = (results_df["TF-IDF Score"] + results_df["BERT Score"]) / 2
        results_df = results_df.sort_values(by="Final Score", ascending=False)
        st.session_state["results_df"] = results_df

        st.write("### Ranked Resumes (Score 1-100)")
        st.dataframe(results_df)
    else:
        st.warning("No valid text extracted from resumes. Please check your files.")

# Search functionality in the main page
if st.session_state.get("show_search", False):
    st.write("### Search Resumes")
    search_query = st.text_input("Enter resume name or keyword:")

    if search_query:
        # Get the ranked results from session state
        results_df = st.session_state.get("results_df", pd.DataFrame())
        
        # Filter results based on the search query
        matched_resumes = results_df[results_df["Resume Name"].str.contains(search_query, case=False, na=False)]
        
        if not matched_resumes.empty:
            st.write("### Matched Resumes")
            # Add a Rank column to the matched results
            matched_resumes["Rank"] = matched_resumes["Final Score"].rank(ascending=False, method="min").astype(int)
            
            # Display the matched results with Rank
            st.dataframe(matched_resumes[["Rank", "Resume Name", "Final Score"]].reset_index(drop=True))

            # Download buttons for matched resumes
            for _, row in matched_resumes.iterrows():
                name = row["Resume Name"]
                if name.endswith(".docx"):
                    pdf_bytes = convert_docx_to_pdf(st.session_state["resume_texts"][name])
                    st.download_button(
                        f"📥 Download {name} as PDF",
                        pdf_bytes,
                        file_name=name.replace(".docx", ".pdf"),
                        mime="application/pdf"
                    )
                else:
                    st.download_button(
                        f"📥 Download {name}",
                        st.session_state["resume_texts"][name],
                        file_name=name,
                        mime="application/pdf"
                    )
        else:
            st.warning("No resumes matched your search query.")
    else:
        st.info("Enter a name or keyword to search.")


# About functionality in the main page
if st.session_state.get("show_about", False):
    st.write("### About")
    st.markdown(
        """
        <div class="about-container">
            <span>About  </span>
            <div class="info-icon">i</div>
        </div>
        """,
        unsafe_allow_html=True
    )
    st.write("""
    This AI-powered system ranks resumes based on their relevance to the job description using NLP and machine learning models.

    ### Features:
    - Extracting text from **PDF and DOCX** resumes.
    - Comparing resumes with **TF-IDF & BERT similarity**.
    - Ranking resumes based on relevance.
    - **Searching** uploaded resumes.
    - **Downloading ranked results**.

    Upload resumes, input a job description, and get ranked recommendations instantly!
    """)